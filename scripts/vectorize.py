import os
import re
from pathlib import Path
# pyrefly: ignore [missing-import]
from dotenv import dotenv_values
from pypdf import PdfReader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
import docx


# 1. 환경 변수 및 API 키 로드
BASE_DIR = Path(__file__).resolve().parent.parent
env_path = BASE_DIR / "ChatbotAPI.env"
env_values = dotenv_values(env_path)
api_key = env_values.get("GEMINI_API_KEY")

if not api_key:
    raise ValueError("ChatbotAPI.env 파일에서 GEMINI_API_KEY를 찾을 수 없습니다.")

os.environ["GOOGLE_API_KEY"] = api_key

# 2. RTF 파서 구현 (사용자가 제공한 .doc 파일이 RTF 형식이므로 맞춤 대응)
sectionchars = {'par': '\n', 'sect': '\n\n', 'page': '\n\n'}
specialchars = {
    'line': '\n', 'tab': '\t', 'emdash': '\u2014', 'endash': '\u2013', 'emspace': '\u2003',
    'enspace': '\u2002', 'qmspace': '\u2005', 'bullet': '\u2022', 'lquote': '\u2018',
    'rquote': '\u2019', 'ldblquote': '\u201c', 'rdblquote': '\u201d', 'row': '\n',
    'cell': ' ', 'nestcell': ' ', '~': '\xa0', '\n': '\n', '\r': '\r', '{': '{', '}': '}',
    '\\': '\\', '-': '\xad', '_': '\u2011',
    **sectionchars
}
destinations = {'fonttbl', 'colortbl', 'stylesheet', 'info', 'listtable', 'listoverridetable'}
PATTERN = re.compile(
    r"\\([a-z]{1,32})(-?\d{1,10})?[ ]?|\\'([0-9a-f]{2})|\\([^a-z])|([{}])|[\r\n]+|(.)",
    re.IGNORECASE,
)

def remove_pict_groups(rtf_text):
    """RTF 내부의 바이너리 이미지 데이터가 텍스트로 오파싱되는 것을 방지하기 위해 제거"""
    if "\\pict" not in rtf_text or "\\bin" not in rtf_text:
        return rtf_text
    result = []
    i = 0
    n = len(rtf_text)
    in_pict = False

    while i < n:
        if not in_pict and rtf_text.startswith("\\pict", i):
            in_pict = True
            # 브레이스 매칭 불균형 해결을 위한 안전장치: 직전 문자가 '{' 라면 삭제
            if result and result[-1] == "{":
                result.pop()
            i += len("\\pict")
            continue

        if in_pict:
            if rtf_text.startswith("\\bin", i):
                i += len("\\bin")
                length_str = ""
                while i < n and rtf_text[i].isdigit():
                    length_str += rtf_text[i]
                    i += 1
                binary_length = int(length_str)
                i += binary_length
                continue
            elif rtf_text[i] == "}":
                in_pict = False
                i += 1
                continue

        if not in_pict:
            result.append(rtf_text[i])

        i += 1

    return "".join(result)

def rtf_to_text(text):
    """RTF 텍스트에서 한글 디코딩 및 제어 문자 제거"""
    text = remove_pict_groups(text)
    
    stack = []
    ignorable = False
    suppress_output = False
    ucskip = 1
    curskip = 0
    hexes = None
    out = ''
    
    for match in PATTERN.finditer(text):
        word, arg, _hex, char, brace, tchar = match.groups()
        
        if hexes and not _hex:
            out += bytes.fromhex(hexes).decode(encoding='cp949', errors='ignore')
            hexes = None
            
        if brace:
            curskip = 0
            if brace == '{':
                stack.append((ucskip, ignorable, suppress_output))
            elif brace == '}':
                if stack:
                    ucskip, ignorable, suppress_output = stack.pop()
                else:
                    ucskip = 0
                    ignorable = False
        elif char:
            curskip = 0
            if char in specialchars:
                if not ignorable:
                    out += specialchars[char]
            elif char == '*':
                ignorable = True
        elif word:
            curskip = 0
            if word in destinations:
                ignorable = True
            if ignorable or suppress_output:
                pass
            elif word in specialchars:
                out += specialchars[word]
            elif word == 'uc':
                ucskip = int(arg)
            elif word == 'u':
                c = int(arg)
                if c < 0:
                    c += 0x10000
                out += chr(c)
                curskip = ucskip
            elif word == 'fonttbl':
                suppress_output = True
            elif word == 'colortbl':
                suppress_output = True
        elif _hex:
            if curskip > 0:
                curskip -= 1
            elif not ignorable:
                if not hexes:
                    hexes = _hex
                else:
                    hexes += _hex
        elif tchar:
            if curskip > 0:
                curskip -= 1
            elif not ignorable and not suppress_output:
                out += tchar
                
    return out

# 3. 개별 파일 로더 함수 정의
def load_document(file_path: Path) -> list[Document]:
    filename = file_path.name
    docs = []
    
    print(f"로딩 중: {filename} ({file_path.stat().st_size:,} bytes)")
    
    # PDF 파일 로드
    if file_path.suffix.lower() == ".pdf":
        reader = PdfReader(file_path)
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": filename, "page": i + 1, "type": "pdf"}
                ))
                
    # TXT 파일 로드 (UTF-8 우선 시도, 실패 시 CP949 시도)
    elif file_path.suffix.lower() == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="cp949") as f:
                content = f.read()
        if content.strip():
            docs.append(Document(
                page_content=content,
                metadata={"source": filename, "type": "txt"}
            ))
            
    # DOCX 파일 로드
    elif file_path.suffix.lower() == ".docx":
        try:
            doc_obj = docx.Document(file_path)
            full_text = []
            # 1. 문단 텍스트 추출
            for para in doc_obj.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            # 2. 표 텍스트 구조적으로 추출
            for table in doc_obj.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    cleaned_row = []
                    for t in row_text:
                        if not cleaned_row or cleaned_row[-1] != t:
                            cleaned_row.append(t)
                    if cleaned_row:
                        full_text.append(" | ".join(cleaned_row))
            content = "\n".join(full_text)
            if content.strip():
                docs.append(Document(
                    page_content=content,
                    metadata={"source": filename, "type": "docx"}
                ))
        except Exception as e:
            print(f"docx 로딩 에러 ({filename}): {e}")
            
    # DOC 파일 (RTF 또는 바이너리 파일 처리)
    elif file_path.suffix.lower() in [".doc", ".rtf"]:
        # RTF 포맷인지 헤더 검증
        try:
            with open(file_path, "r", encoding="latin1") as f:
                header = f.read(10)
        except Exception:
            header = ""
            
        if header.startswith("{\\rtf1"):
            with open(file_path, "r", encoding="latin1") as f:
                content = f.read()
            text = rtf_to_text(content)
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": filename, "type": "rtf"}
                ))
        else:
            print(f"경고: {filename} 파일은 바이너리 Word 파일이거나 RTF 헤더가 불일치하여 현재 파싱을 건너뜁니다.")
            
    return docs

def main():
    data_dir = BASE_DIR / "RAG 백데이터"
    db_dir = BASE_DIR / "chromadb_store"
    
    if not data_dir.exists():
        print(f"오류: '{data_dir}' 폴더가 존재하지 않습니다.")
        return
        
    # 기존 DB 디렉토리가 있다면 안전하게 제거하여 초기화
    if db_dir.exists():
        print(f"기존 Chroma DB 초기화 중... ({db_dir})")
        import shutil
        try:
            shutil.rmtree(db_dir)
            print("기존 DB 디렉토리 삭제 완료.")
        except Exception as e:
            print(f"기존 DB 디렉토리 삭제 실패 ({e}). Chroma 내부 문서 삭제를 시도합니다.")
            try:
                temp_db = Chroma(
                    persist_directory=str(db_dir),
                    embedding_function=embeddings
                )
                ids = temp_db.get()['ids']
                if ids:
                    temp_db.delete(ids)
                    print(f"Chroma 내부 문서 {len(ids)}개 일괄 삭제 완료.")
            except Exception as ex:
                print(f"Chroma 내부 문서 삭제 실패: {ex}.")
        
    all_raw_docs = []
    
    # 폴더 내 모든 파일 탐색
    for file_path in data_dir.iterdir():
        if file_path.is_file():
            try:
                loaded_docs = load_document(file_path)
                all_raw_docs.extend(loaded_docs)
            except Exception as e:
                print(f"에러 발생 ({file_path.name}): {e}")
                
    if not all_raw_docs:
        print("경고: 파싱된 데이터가 존재하지 않습니다.")
        return
        
    print(f"\n총 {len(all_raw_docs)}개의 페이지/문서를 읽어왔습니다. 텍스트 분할을 시작합니다.")
    
    # 4. 텍스트 분할 (Chunking)
    # 한 조항 단위 또는 문맥이 끊기지 않는 크기로 쪼갬
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,
        chunk_overlap=300,
        length_function=len
    )
    chunks = text_splitter.split_documents(all_raw_docs)
    print(f"분할 결과: 총 {len(chunks)}개의 청크(Chunk)로 나뉘었습니다.")
    
    # 5. 임베딩 모델 및 Chroma DB 구성
    print("\nGoogle Gemini 임베딩 모델을 사용하여 벡터화를 진행합니다. (models/gemini-embedding-2)")
    embeddings = GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-2")
    
    print(f"Chroma DB에 임베딩 데이터를 저장합니다... (저장 위치: {db_dir})")
    
    import time
    
    # 429 Rate Limit 발생 시 지수 백오프(Exponential Backoff)를 이용해 자동 재시도하는 함수 정의
    def add_documents_with_retry(vectordb, documents, batch_size=20, base_sleep=20):
        for i in range(0, len(documents), batch_size):
            batch = documents[i:i+batch_size]
            success = False
            retry_delay = base_sleep
            
            while not success:
                try:
                    print(f"임베딩 적재 중... {i}/{len(documents)}")
                    if i == 0 and vectordb is None:
                        # 첫 배치로 벡터 DB를 생성하는 경우
                        vectordb = Chroma.from_documents(
                            documents=batch,
                            embedding=embeddings,
                            persist_directory=str(db_dir)
                        )
                    else:
                        vectordb.add_documents(batch)
                    success = True
                    # 성공적인 호출 후 안전을 위해 약간 대기
                    time.sleep(3)
                except Exception as e:
                    print(f"\n[오류 발생] {e}")
                    print(f"네트워크 및 API 할당량 우회를 위해 {retry_delay}초 대기 후 다시 시도합니다...")
                    time.sleep(retry_delay)
                    retry_delay = min(retry_delay * 2, 120)  # 최대 2분까지 대기 시간 확장
        return vectordb

    # 배치 크기를 20으로 줄여 단일 요청의 크기를 줄임
    vectordb = add_documents_with_retry(None, chunks, batch_size=20, base_sleep=20)
    
    print("\n성공적으로 RAG 백데이터 벡터베이스를 구축하였습니다!")
    print(f"생성된 벡터 DB 총 아이템 수: {vectordb._collection.count()}")

if __name__ == "__main__":
    main()
