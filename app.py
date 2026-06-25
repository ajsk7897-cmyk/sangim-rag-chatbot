import os
os.environ["PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION"] = "python"

try:
    __import__('pysqlite3')
    import sys
    sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
except ImportError:
    pass

import streamlit as st
from pathlib import Path
from dotenv import load_dotenv
from langchain_chroma import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.documents import Document
from pypdf import PdfReader
import docx
import shutil
import time
import re


# 1. 페이지 초기 설정 및 테마 로드 (Premium UI CSS 주입)
st.set_page_config(
    page_title="상임법 RAG 챗봇",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom Premium CSS 주입
st.markdown("""
<style>
    /* Google Font Outfit */
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600&family=Noto+Sans+KR:wght@300;400;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Noto Sans KR', 'Outfit', sans-serif;
    }
    
    /* Main Background & Glassmorphism Sidebar */
    .stApp {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        color: #f8fafc;
    }
    
    [data-testid="stSidebar"] {
        background: rgba(30, 41, 59, 0.7) !important;
        backdrop-filter: blur(10px);
        border-right: 1px solid rgba(255, 255, 255, 0.05);
    }
    
    /* Header Styles */
    h1 {
        background: linear-gradient(90deg, #38bdf8 0%, #818cf8 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 700;
        font-size: 2.2rem;
        margin-bottom: 0.5rem;
    }
    
    /* Chat Bubble Custom Styling */
    .stChatMessage {
        border-radius: 16px !important;
        padding: 1rem !important;
        margin-bottom: 0.8rem !important;
        border: 1px solid rgba(255, 255, 255, 0.05);
        color: #ffffff !important;
    }
    
    .stChatMessage p, .stChatMessage span, .stChatMessage div, .stChatMessage li {
        color: #ffffff !important;
    }
    
    /* User Chat Bubble */
    [data-testid="stChatMessage"]:nth-child(even) {
        background-color: rgba(56, 189, 248, 0.1) !important;
        border-left: 4px solid #38bdf8 !important;
    }
    
    /* Assistant Chat Bubble */
    [data-testid="stChatMessage"]:nth-child(odd) {
        background-color: rgba(30, 41, 59, 0.6) !important;
        border-left: 4px solid #818cf8 !important;
    }
    
    /* Hover Effects for interactive elements */
    .stButton>button {
        background: linear-gradient(90deg, #0284c7 0%, #4f46e5 100%) !important;
        color: white !important;
        border: none !important;
        border-radius: 8px !important;
        padding: 0.5rem 1.5rem !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
    }
    .stButton>button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 4px 12px rgba(56, 189, 248, 0.3) !important;
    }
    
    /* Source box styling */
    .source-box {
        background-color: rgba(15, 23, 42, 0.5);
        border: 1px solid rgba(255, 255, 255, 0.1);
        border-radius: 8px;
        padding: 0.8rem;
        margin-top: 0.5rem;
        font-size: 0.9rem;
    }
    
    /* Smooth Scroll */
    html {
        scroll-behavior: smooth;
    }
</style>
""", unsafe_allow_html=True)

# 2. API 키 로드
BASE_DIR = Path(__file__).resolve().parent

# 1) OS 환경변수 또는 Streamlit Secrets에서 우선 탐색 (배포 환경용)
api_key = os.environ.get("GEMINI_API_KEY")

# st.secrets 지원 여부 확인 후 로드
if not api_key:
    try:
        api_key = st.secrets.get("GEMINI_API_KEY")
    except Exception:
        pass

# 2) 찾지 못한 경우 로컬 ChatbotAPI.env 파일에서 로드 (로컬 개발용)
if not api_key:
    env_path = BASE_DIR / "ChatbotAPI.env"
    if env_path.exists():
        load_dotenv(env_path)
        api_key = os.environ.get("GEMINI_API_KEY")

if not api_key:
    st.error("🔑 API 키를 찾을 수 없습니다. 로컬의 경우 ChatbotAPI.env 파일을 확인해 주시고, 클라우드 배포인 경우 Secrets 설정을 확인해 주세요.")
    st.stop()

os.environ["GOOGLE_API_KEY"] = api_key

# --- RTF 파싱 도구 ---
sectionchars = {'par': '\n', 'sect': '\n\n', 'page': '\n\n'}
specialchars = {
    'line': '\n', 'tab': '\t', 'emdash': '\u2014', 'endash': '\u2013', 'emspace': '\u2003',
    'enspace': '\u2002', 'qmspace': '\u2005', 'bullet': '\u2022', 'lquote': '\u2018',
    'rquote': '\u2019', 'ldblquote': '\u201c', 'rdblquote': '\u201d', 'row': '\n',
    'cell': ' ', 'nestcell': ' ', '~': '\xa0', '\n': '\n', '\r': '\r', '{': '{', '}': '}',
    '\\': '\\', '-': '\xad', '_': '\u2011',
    **sectionchars
}
destinations = {'fonttbl', 'colortbl', 'stylesheet', 'info', 'listtable', 'listoverridetable'}
PATTERN = re.compile(
    r"\\([a-z]{1,32})(-?\d{1,10})?[ ]?|\\'([0-9a-f]{2})|\\([^a-z])|([{}])|[\r\n]+|(.)",
    re.IGNORECASE,
)

def remove_pict_groups(rtf_text):
    if "\\pict" not in rtf_text or "\\bin" not in rtf_text:
        return rtf_text
    result = []
    i = 0
    n = len(rtf_text)
    in_pict = False
    while i < n:
        if not in_pict and rtf_text.startswith("\\pict", i):
            in_pict = True
            if result and result[-1] == "{":
                result.pop()
            i += len("\\pict")
            continue
        if in_pict:
            if rtf_text.startswith("\\bin", i):
                i += len("\\bin")
                length_str = ""
                while i < n and rtf_text[i].isdigit():
                    length_str += rtf_text[i]
                    i += 1
                binary_length = int(length_str)
                i += binary_length
                continue
            elif rtf_text[i] == "}":
                in_pict = False
                i += 1
                continue
        if not in_pict:
            result.append(rtf_text[i])
        i += 1
    return "".join(result)

def rtf_to_text(text):
    text = remove_pict_groups(text)
    stack = []
    ignorable = False
    suppress_output = False
    ucskip = 1
    curskip = 0
    hexes = None
    out = ''
    for match in PATTERN.finditer(text):
        word, arg, _hex, char, brace, tchar = match.groups()
        if hexes and not _hex:
            out += bytes.fromhex(hexes).decode(encoding='cp949', errors='ignore')
            hexes = None
        if brace:
            curskip = 0
            if brace == '{':
                stack.append((ucskip, ignorable, suppress_output))
            elif brace == '}':
                if stack:
                    ucskip, ignorable, suppress_output = stack.pop()
                else:
                    ucskip = 0
                    ignorable = False
        elif char:
            curskip = 0
            if char in specialchars:
                if not ignorable:
                    out += specialchars[char]
            elif char == '*':
                ignorable = True
        elif word:
            curskip = 0
            if word in destinations:
                ignorable = True
            if ignorable or suppress_output:
                pass
            elif word in specialchars:
                out += specialchars[word]
            elif word == 'uc':
                ucskip = int(arg)
            elif word == 'u':
                c = int(arg)
                if c < 0:
                    c += 0x10000
                out += chr(c)
                curskip = ucskip
            elif word == 'fonttbl':
                suppress_output = True
            elif word == 'colortbl':
                suppress_output = True
        elif _hex:
            if curskip > 0:
                curskip -= 1
            elif not ignorable:
                if not hexes:
                    hexes = _hex
                else:
                    hexes += _hex
        elif tchar:
            if curskip > 0:
                curskip -= 1
            elif not ignorable and not suppress_output:
                out += tchar
    return out

# 3. 데이터베이스 및 RAG 모델 로드 함수 (캐싱 적용)
@st.cache_resource
def load_rag_system():
    db_dir = BASE_DIR / "chromadb_store"
    if not db_dir.exists():
        return None, None
    
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-2",
        google_api_key=api_key
    )
    vectordb = Chroma(
        persist_directory=str(db_dir),
        embedding_function=embeddings
    )
    
    # 챗봇용 LLM 로드 (Gemini 1.5 Flash 사용 - 무료 티어 API 한도 내에서 최고의 성능)
    llm = ChatGoogleGenerativeAI(
        model="gemini-1.5-flash",
        temperature=0.1,  # 정밀한 사실 기반 답변을 위해 온도를 낮춤
        max_output_tokens=4096,
        google_api_key=api_key,
        max_retries=6
    )
    
    return vectordb, llm

vectordb, llm = load_rag_system()

# 4. 실시간 백데이터 동기화 및 벡터화 실행 함수
def sync_and_vectorize():
    from langchain_text_splitters import RecursiveCharacterTextSplitter
    data_dir = BASE_DIR / "RAG 백데이터"
    db_dir = BASE_DIR / "chromadb_store"
    
    if not data_dir.exists():
        data_dir.mkdir(parents=True, exist_ok=True)
        
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-2",
        google_api_key=api_key
    )
    # 기존 DB 디렉토리가 있다면 안전하게 제거하여 초기화
    if db_dir.exists():
        try:
            shutil.rmtree(db_dir)
        except Exception as e:
            st.warning(f"기존 DB 디렉토리 삭제 실패 ({e}). Chroma 내부 문서 일괄 삭제를 시도합니다.")
            try:
                temp_db = Chroma(
                    persist_directory=str(db_dir),
                    embedding_function=embeddings
                )
                ids = temp_db.get()['ids']
                if ids:
                    temp_db.delete(ids)
            except Exception as ex:
                st.error(f"Chroma 내부 문서 삭제 실패: {ex}")
            
    all_raw_docs = []
    
    # 폴더 내 모든 파일 탐색
    for file_path in data_dir.iterdir():
        if file_path.is_file():
            filename = file_path.name
            suffix = file_path.suffix.lower()
            
            # PDF 파일 로드
            if suffix == ".pdf":
                try:
                    reader = PdfReader(file_path)
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text() or ""
                        if text.strip():
                            all_raw_docs.append(Document(
                                page_content=text,
                                metadata={"source": filename, "page": i + 1, "type": "pdf"}
                            ))
                except Exception as e:
                    st.warning(f"PDF 로딩 에러 ({filename}): {e}")
                    
            # TXT 파일 로드 (UTF-8 우선 시도, 실패 시 CP949 시도)
            elif suffix == ".txt":
                try:
                    try:
                        with open(file_path, "r", encoding="utf-8") as f:
                            content = f.read()
                    except UnicodeDecodeError:
                        with open(file_path, "r", encoding="cp949") as f:
                            content = f.read()
                    if content.strip():
                        all_raw_docs.append(Document(
                            page_content=content,
                            metadata={"source": filename, "type": "txt"}
                        ))
                except Exception as e:
                    st.warning(f"TXT 로딩 에러 ({filename}): {e}")
                    
            # DOCX 파일 로드
            elif suffix == ".docx":
                try:
                    doc_obj = docx.Document(file_path)
                    full_text = []
                    # 1. 문단 텍스트 추출
                    for para in doc_obj.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)
                    # 2. 표 텍스트 구조적으로 추출
                    for table in doc_obj.tables:
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            cleaned_row = []
                            for t in row_text:
                                if not cleaned_row or cleaned_row[-1] != t:
                                    cleaned_row.append(t)
                            if cleaned_row:
                                full_text.append(" | ".join(cleaned_row))
                    content = "\n".join(full_text)
                    if content.strip():
                        all_raw_docs.append(Document(
                            page_content=content,
                            metadata={"source": filename, "type": "docx"}
                        ))
                except Exception as e:
                    st.warning(f"docx 로딩 에러 ({filename}): {e}")
                    
            # DOC/RTF 파일 로드
            elif suffix in [".doc", ".rtf"]:
                try:
                    with open(file_path, "r", encoding="latin1") as f:
                        header = f.read(10)
                except Exception:
                    header = ""
                    
                if header.startswith("{\\rtf1"):
                    try:
                        with open(file_path, "r", encoding="latin1") as f:
                            content = f.read()
                        text = rtf_to_text(content)
                        if text.strip():
                            all_raw_docs.append(Document(
                                page_content=text,
                                metadata={"source": filename, "type": "rtf"}
                            ))
                    except Exception as e:
                        st.warning(f"RTF 로딩 에러 ({filename}): {e}")
                        
    if not all_raw_docs:
        return False, "파싱된 문서가 존재하지 않습니다."
        
    # 텍스트 분할 (Chunking) - 최적화 설정
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1200,
        chunk_overlap=300,
        length_function=len
    )
    chunks = text_splitter.split_documents(all_raw_docs)
    
    # 임베딩 모델 및 Chroma DB 구성 (함수 초입에 정의된 embeddings 재사용)
    
    batch_size = 20
    base_sleep = 20
    vectordb = None
    
    progress_bar = st.progress(0.0)
    status_text = st.empty()
    
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i+batch_size]
        success = False
        retry_delay = base_sleep
        
        progress_val = min(i / len(chunks), 1.0)
        progress_bar.progress(progress_val)
        status_text.text(f"임베딩 적재 중... ({i}/{len(chunks)} 청크 완료)")
        
        while not success:
            try:
                if i == 0:
                    vectordb = Chroma.from_documents(
                        documents=batch,
                        embedding=embeddings,
                        persist_directory=str(db_dir)
                    )
                else:
                    vectordb.add_documents(batch)
                success = True
                time.sleep(3)
            except Exception as e:
                status_text.text(f"API 할당량 초과 등으로 재시도 대기 중 ({retry_delay}초): {e}")
                time.sleep(retry_delay)
                retry_delay = min(retry_delay * 2, 120)
                
    progress_bar.progress(1.0)
    status_text.text(f"동기화 적재 완료! 총 {len(chunks)}개 청크 저장됨.")
    return True, len(chunks)


# 사이드바 구성
with st.sidebar:
    st.title("⚙️ 설정 및 도구")
    st.markdown("---")
    
    # DB 상태 표시
    if vectordb is not None:
        db_count = vectordb._collection.count()
        st.success(f"데이터베이스 연결 성공\n(총 {db_count}개 청크 적재됨)")
    else:
        st.warning("⚠️ 백데이터 데이터베이스(chromadb_store)가 존재하지 않습니다. 먼저 `python scripts/vectorize.py`를 실행해 주세요.")
        
    st.markdown("---")
    
    # RAG 검색 설정
    k_value = st.slider("참조할 단락 개수 (K)", min_value=2, max_value=8, value=4, step=1)
    show_sources = st.checkbox("답변 시 참조 출처 표시", value=True)
    
    st.markdown("---")
    # 대화 초기화 버튼
    if st.button("🔄 대화 기록 초기화"):
        st.session_state.messages = []
        st.rerun()
        
    st.markdown("---")
    st.markdown("<p style='font-size: 0.8rem; color: #64748b;'>상가건물 임대차보호법 RAG 챗봇 v1.0</p>", unsafe_allow_html=True)

# 메인 헤더
st.title("⚖️ 상임법 전문 질의응답 챗봇")
st.markdown("내가 제공한 백데이터(상가임대차 보호법 전문, 상담 사례집, 주요 판례 등)에 근거하여 관련 법률 질문에 답변해 드립니다.")
st.markdown("---")

tab1, tab2 = st.tabs(["💬 챗봇 대화", "📂 백데이터 관리"])

with tab1:
    # 대화 세션 상태 초기화
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    # 대화 기록 렌더링
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
            if "sources" in message and message["sources"] and show_sources:
                with st.expander("📚 참조된 백데이터 확인"):
                    for idx, src in enumerate(message["sources"]):
                        st.markdown(f"**[{idx+1}] {src['source']}** (유사도 스코어: {src['score']:.4f})")
                        st.caption(src["content"])
                        
    # 사용자 입력 처리
    if question := st.chat_input("상가건물 임대차보호법에 대해 질문해 보세요 (예: 임차인의 권리금 회수 기회 보장 기간은?)"):
        # 사용자 질문 화면 렌더링 및 저장
        with st.chat_message("user"):
            st.markdown(question)
        st.session_state.messages.append({"role": "user", "content": question})
        
        # RAG 답변 생성
        with st.chat_message("assistant"):
            response_placeholder = st.empty()
            
            if vectordb is None or llm is None:
                ans = "오류: 데이터베이스가 정상적으로 설정되지 않았습니다. 관리자에게 문의하세요."
                response_placeholder.markdown(ans)
                st.session_state.messages.append({"role": "assistant", "content": ans})
            else:
                # 1. 유사한 법률 데이터 검색
                docs_with_scores = None
                try:
                    docs_with_scores = vectordb.similarity_search_with_relevance_scores(question, k=k_value)
                except Exception as e:
                    error_msg = str(e)
                    if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                        delay_match = re.search(r"retry in (\d+\.?\d*)s", error_msg)
                        if delay_match:
                            delay_sec = float(delay_match.group(1))
                            ans = f"⚠️ 임베딩 API 무료 사용량 한도를 초과했습니다. 약 {int(delay_sec) + 1}초 후에 다시 질문해 주세요."
                        else:
                            ans = "⚠️ 임베딩 API 무료 사용량 한도를 초과했습니다. 잠시 후(약 1분) 다시 시도해 주세요."
                    else:
                        ans = f"⚠️ 문서 검색 중 오류가 발생했습니다: {e}"
                    response_placeholder.markdown(ans)
                    st.session_state.messages.append({"role": "assistant", "content": ans})
                    st.rerun()
                
                # 2. Context 구성 및 출처 추출
                context_list = []
                sources_info = []
                
                for doc, score in docs_with_scores:
                    context_list.append(doc.page_content)
                    sources_info.append({
                        "source": doc.metadata.get("source", "알 수 없음"),
                        "page": doc.metadata.get("page", None),
                        "content": doc.page_content,
                        "score": score
                    })
                    
                context = "\n\n".join(context_list)
                
                # 3. 시스템 프롬프트 작성 (백데이터 중심 제약조건 부여 및 상황 해석 기능 추가)
                system_prompt = """당신은 상가건물 임대차보호법(상임법) 전문 상담 챗봇입니다.
반드시 아래에 주어진 법률 데이터(Context)만을 근거로 삼아 사용자의 질문에 답변하십시오.

[핵심 제약 조건 및 해석 가이드]
1. 제공된 법률 데이터(Context)에 직접적으로 언급된 내용이거나, 이를 바탕으로 논리적으로 유추·적용할 수 있는 내용만 사실에 입각하여 답변하십시오.
2. 사용자가 구체적인 상황(예: 계약기간 10년 경과 후 퇴거 요구 등)을 제시할 경우, 제공된 판례나 법 조항의 원칙을 해당 상황에 논리적으로 대입하여 해석한 결과를 친절하게 설명해 주십시오.
3. 제공된 데이터만으로 사용자의 상황에 대해 확실하게 판단할 수 없거나 관련 내용이 전혀 존재하지 않는 경우, 절대로 임의로 지어내거나 외부 인터넷 검색, 사전 훈련된 일반 상식으로 답변하지 마십시오. 이때는 "제공된 백데이터에서 관련 정보를 찾을 수 없습니다."라고 답변하십시오.
4. 법률 조항이나 판례 등을 언급할 경우 구체적인 조항번호나 출처를 명시하십시오.
5. [개념 혼동 방지]: '상임법 전면 적용 대상 환산보증금 상한액'과 '최우선변제를 받기 위한 소액임차인 범위 환산보증금 기준'은 완전히 다른 제도적 개념입니다. 절대 두 기준과 금액을 섞거나 잘못 매칭하지 말고 명확하게 분리하여 설명하십시오.
6. [면책 조항]: 당신이 제공하는 해석은 AI의 논리적 추론일 뿐 법적 효력이 없음을 반드시 답변의 마지막에 안내하십시오. (예: "본 답변은 제공된 데이터를 바탕으로 한 AI의 해석이므로 법적 효력이 없으며, 실제 분쟁 시에는 전문가의 상담을 권장합니다.")

[답변 예시]
질문: 계약기간 10년이 경과한 임차인이 임대인의 재계약 조건을 받아들이지 않을 시 퇴거가 가능해?
답변: 제공된 판례 및 상가건물 임대차보호법에 따르면, 임차인의 계약갱신요구권은 최초의 임대차기간을 포함한 전체 임대차기간이 10년을 초과하지 아니하는 범위에서만 행사할 수 있습니다(제10조 제2항).
따라서 전체 임대차기간이 10년을 경과한 경우 임차인은 더 이상 갱신요구권을 행사할 수 없으므로, 임대인과 새로운 조건으로 합의(재계약)하지 못한다면 임대인의 명도(퇴거) 요구에 응해야 할 가능성이 높습니다.
(참고 판례: 대법원 20XX.X.X. 선고 20XX다XXXX 판결)

⚠️ 본 답변은 제공된 백데이터를 바탕으로 한 AI의 해석이므로 법적 효력이 없으며, 실제 분쟁 시에는 변호사 등 전문가의 상담을 권장합니다.

제공된 법률 데이터 (Context):
{context}"""

                # 4. LLM 체인 호출
                prompt = ChatPromptTemplate.from_messages([
                    ("system", system_prompt),
                    ("human", "{question}")
                ])
                
                chain = prompt | llm
                
                with st.spinner("백데이터에서 조문을 조회하고 답변을 준비하는 중입니다..."):
                    try:
                        stream_res = chain.stream({"context": context, "question": question})
                    except Exception as e:
                        error_msg = str(e)
                        if "429" in error_msg or "RESOURCE_EXHAUSTED" in error_msg:
                            delay_match = re.search(r"retry in (\d+\.?\d*)s", error_msg)
                            if delay_match:
                                delay_sec = float(delay_match.group(1))
                                st.error(f"⚠️ 답변 생성 API 무료 사용량 한도를 초과했습니다. 약 {int(delay_sec) + 1}초 후에 다시 질문해 주세요.")
                            else:
                                st.error("⚠️ 답변 생성 API 무료 사용량 한도를 초과했습니다. 잠시 후 다시 시도해 주세요.")
                        else:
                            st.error(f"LLM 호출 중 에러가 발생했습니다: {e}")
                        stream_res = None
                        
                # 5. 최종 답변 실시간 스트리밍 출력 (st.write_stream 활용)
                answer = ""
                if stream_res:
                    # 기존에 생성된 placeholder를 비우고 native write_stream 사용
                    response_placeholder.empty()
                    try:
                        def generate_chunks():
                            for chunk in stream_res:
                                content = chunk.content
                                if isinstance(content, str):
                                    yield content
                                elif isinstance(content, list):
                                    for item in content:
                                        if isinstance(item, str):
                                            yield item
                                        elif isinstance(item, dict) and "text" in item:
                                            yield item["text"]
                                            
                        # st.write_stream을 사용하여 타자 치듯 실시간 출력
                        answer = st.write_stream(generate_chunks())
                    except Exception as e:
                        st.warning(f"\n[스트리밍 중간 끊김 발생] 일부 답변만 로드되었습니다: {e}")
                        
                if not answer:
                    answer = "답변 생성에 실패했습니다. (API 오류 또는 제한)"
                    st.markdown(answer)
                
                # 6. 소스 및 답변 세션 저장
                st.session_state.messages.append({
                    "role": "assistant",
                    "content": answer,
                    "sources": sources_info
                })
                
                # 7. 출처 박스 렌더링
                if sources_info and show_sources:
                    with st.expander("📚 참조된 백데이터 확인"):
                        for idx, src in enumerate(sources_info):
                            page_str = f" p.{src['page']}" if src['page'] else ""
                            st.markdown(f"**[{idx+1}] {src['source']}{page_str}** (유사도 스코어: {src['score']:.4f})")
                            st.caption(src["content"])
                            
                # 화면 리런으로 상태 동기화
                st.rerun()

with tab2:
    st.subheader("📂 RAG 백데이터 관리")
    st.markdown("챗봇이 참고할 원본 문서를 추가하거나 삭제할 수 있습니다. 파일을 관리한 후 반드시 **하단의 동기화 버튼**을 눌러주셔야 데이터베이스가 업데이트됩니다.")
    st.markdown("---")
    
    # 데이터 폴더 정의 및 생성
    data_dir = BASE_DIR / "RAG 백데이터"
    if not data_dir.exists():
        data_dir.mkdir(parents=True, exist_ok=True)
        
    # 1. 파일 업로드 섹션
    st.markdown("### 📤 신규 백데이터 파일 업로드")
    uploaded_files = st.file_uploader(
        "지원 포맷: .pdf, .docx, .txt, .rtf (여러 파일 업로드 가능)",
        type=["pdf", "docx", "txt", "rtf", "doc"],
        accept_multiple_files=True
    )
    
    if uploaded_files:
        save_btn = st.button("💾 업로드한 파일 저장하기", use_container_width=True)
        if save_btn:
            for uploaded_file in uploaded_files:
                target_path = data_dir / uploaded_file.name
                with open(target_path, "wb") as f:
                    f.write(uploaded_file.getbuffer())
            st.success(f"{len(uploaded_files)}개의 파일이 성공적으로 저장되었습니다.")
            time.sleep(1)
            st.rerun()
            
    st.markdown("---")
    
    # 2. 파일 목록 조회 및 제거 섹션
    st.markdown("### 📋 현재 적재된 원본 백데이터 파일 목록")
    files = [f for f in data_dir.iterdir() if f.is_file()]
    
    if not files:
        st.info("현재 저장된 백데이터 파일이 없습니다. 문서를 업로드해 주세요.")
    else:
        for idx, file in enumerate(files):
            col1, col2, col3 = st.columns([6, 2, 2])
            size_kb = file.stat().st_size / 1024
            ext = file.suffix.upper()
            
            col1.markdown(f"**{idx+1}. {file.name}**")
            col2.caption(f"{ext} 파일 / {size_kb:.1f} KB")
            
            # 삭제 버튼 클릭 핸들링
            if col3.button(f"❌ 삭제", key=f"del_{file.name}"):
                try:
                    os.remove(file)
                    st.success(f"{file.name} 파일이 삭제되었습니다.")
                    time.sleep(1)
                    st.rerun()
                except Exception as e:
                    st.error(f"파일 삭제 오류: {e}")
                    
    st.markdown("---")
    
    # 3. 외부 데이터 자동 수집 기능
    st.markdown("### 🌐 외부 데이터(API/뉴스레터) 자동 수집")
    st.markdown("국가법령정보센터 판례 및 최신 웹 기사/뉴스레터를 자동으로 수집하여 백데이터로 저장합니다.")
    
    col_api, col_crawl = st.columns(2)
    api_btn = col_api.button("⚖️ 대법원 판례 수집 (Open API)", use_container_width=True)
    crawl_btn = col_crawl.button("📰 최신 뉴스/칼럼 수집 (크롤링)", use_container_width=True)
    
    if api_btn:
        with st.spinner("국가법령정보센터에서 판례를 수집하고 있습니다..."):
            import subprocess
            try:
                result = subprocess.run(["python", str(BASE_DIR / "scripts" / "collect_law_api.py")], capture_output=True, text=True, encoding="utf-8")
                if result.returncode == 0:
                    st.success("대법원 판례 수집 완료!")
                    st.code(result.stdout, language="text")
                else:
                    st.error(f"수집 실패: {result.stderr}")
            except Exception as e:
                st.error(f"실행 오류: {e}")
                
    if crawl_btn:
        with st.spinner("웹에서 최신 칼럼과 뉴스를 수집하고 있습니다..."):
            import subprocess
            try:
                result = subprocess.run(["python", str(BASE_DIR / "scripts" / "collect_crawling.py")], capture_output=True, text=True, encoding="utf-8")
                if result.returncode == 0:
                    st.success("뉴스 및 칼럼 수집 완료!")
                    st.code(result.stdout, language="text")
                else:
                    st.error(f"수집 실패: {result.stderr}")
            except Exception as e:
                st.error(f"실행 오류: {e}")

    st.markdown("---")
    
    # 4. DB 동기화 실행 버튼
    st.markdown("### ⚡ 데이터베이스 동기화")
    st.warning("⚠️ 백데이터 파일을 추가하거나 삭제한 후 아래 버튼을 클릭하여 Chroma DB에 임베딩을 다시 구워야 챗봇이 최신 데이터를 기반으로 대답합니다. (기존 DB는 초기화 후 새로 구축됩니다.)")
    
    sync_btn = st.button("⚡ 백데이터 동기화 및 DB 업데이트 실행", use_container_width=True)
    if sync_btn:
        with st.spinner("백데이터 문서를 분석하고 벡터 DB에 저장하는 중입니다... (API 할당량 규정 준수를 위해 다소 시간이 소요될 수 있습니다)"):
            success, result = sync_and_vectorize()
            if success:
                st.success(f"🎉 동기화 성공! 총 {result}개의 청크가 데이터베이스에 성공적으로 적재되었습니다.")
                st.cache_resource.clear() # streamlit 캐시 해제하여 챗봇 로딩 시 최신 DB 로드 유도
                time.sleep(2)
                st.rerun()
            else:
                st.error(f"동기화 실패: {result}")

